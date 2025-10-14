from fastapi import UploadFile
import io
import pandas as pd
from docx import Document
import pdfplumber
from typing import Dict, Any
import logging

logger = logging.getLogger(__name__)


async def extract_text_from_file(file: UploadFile) -> Dict[str, Any]:
    """Извлекает текст и структурированные данные из файла"""
    content = await file.read()
    filename = file.filename.lower()

    result = {
        "text": "",
        "tables": [],
        "metadata": {"filename": file.filename, "type": filename.split('.')[-1]}
    }

    try:
        if filename.endswith(".txt"):
            result["text"] = content.decode("utf-8", errors="ignore")

        elif filename.endswith(".docx"):
            doc = Document(io.BytesIO(content))
            result["text"] = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

            # Извлекаем таблицы из DOCX
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    result["tables"].append(table_data)

        elif filename.endswith(".pdf"):
            text = ""
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    text += page_text + "\n"

                    # Извлекаем таблицы из PDF
                    tables = page.extract_tables()
                    for table in tables:
                        if table and any(any(cell is not None for cell in row) for row in table):
                            result["tables"].append(table)

            result["text"] = text.strip()

        elif filename.endswith(".xlsx"):
            # Обработка Excel файлов
            excel_file = io.BytesIO(content)

            # Читаем все листы
            xl = pd.ExcelFile(excel_file)
            all_text = f"Excel файл: {file.filename}\n"
            all_tables = []

            for sheet_name in xl.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                all_text += f"\n--- Лист: {sheet_name} ---\n"
                all_text += df.to_string() + "\n"

                # Сохраняем таблицу как список списков
                all_tables.append({
                    "sheet_name": sheet_name,
                    "data": df.fillna("").values.tolist(),
                    "columns": df.columns.tolist()
                })

            result["text"] = all_text
            result["tables"] = all_tables

        else:
            raise ValueError(f"Неподдерживаемый формат файла: {filename}")

        logger.info(
            f"Успешно обработан файл {file.filename}: {len(result['text'])} символов, {len(result['tables'])} таблиц")
        return result

    except Exception as e:
        logger.error(f"Ошибка обработки файла {file.filename}: {e}")
        raise ValueError(f"Ошибка обработки файла: {str(e)}")

